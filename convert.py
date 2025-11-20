import os
import tempfile
import pypandoc
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def convert_markdowns_to_docx(md_files=None, output_file="combined.docx", chapters_file=None, folder_name=None, reference_docx=None):
    """
    Convert markdown files to DOCX using Pandoc, handling <partname> logic from chapters.txt.
    Args:
        md_files: List of markdown files (used if chapters_file is None)
        output_file: Output DOCX file path
        chapters_file: Path to chapters.txt file (takes precedence over md_files)
        folder_name: Optional folder where markdown files are located
        reference_docx: Optional DOCX template for styling
    """
    temp_files = []
    combined_md = ""
    if chapters_file and os.path.exists(chapters_file):
        with open(chapters_file, encoding="utf-8") as f:
            lines = [line.strip() for line in f if line.strip() and not line.strip().startswith("#")]
        for line in lines:
            if line.startswith("<partname>"):
                part_title = line.replace("<partname>", "").strip()
                combined_md += f"\n\n# {part_title}\n\n<!-- PAGEBREAK -->\n\n"
            else:
                chapter_file = line
                if folder_name and not os.path.isabs(chapter_file):
                    chapter_file = os.path.join(folder_name, chapter_file)
                if not chapter_file.endswith(".md"):
                    chapter_file += ".md"
                if os.path.exists(chapter_file):
                    with open(chapter_file, encoding="utf-8") as chf:
                        combined_md += chf.read() + "\n\n<!-- PAGEBREAK -->\n\n"
    elif md_files:
        for f in md_files:
            chapter_file = os.path.join(folder_name, f) if folder_name and not os.path.isabs(f) else f
            if os.path.exists(chapter_file):
                with open(chapter_file, encoding="utf-8") as chf:
                    combined_md += chf.read() + "\n\n:::pagebreak:::\n\n"
    else:
        raise ValueError("Either chapters_file must exist or md_files must be provided")

    # Suppress alt text captions for images by removing alt text from markdown
    def suppress_image_alt(md):
        # Replace ![alt](src) with ![](src)
        return re.sub(r'!\[[^\]]*\]\(([^\)]+)\)', r'![&nbsp;](\1)', md)
    def reset_caps_in_code(md):
        return re.sub(r'(\n```[a-zA-Z0-9_+-]*)', r'\1', md)
    combined_md = suppress_image_alt(combined_md)
    combined_md = reset_caps_in_code(combined_md)

    # Write combined markdown to a temp file
    temp_md_file = tempfile.NamedTemporaryFile(delete=False, suffix=".md", mode="w", encoding="utf-8")
    temp_md_file.write(combined_md)
    temp_md_file.close()
    temp_files.append(temp_md_file.name)

    # Helper to create a temporary reference DOCX that contains the requested
    # centered header and a centered page-number footer. If `reference_docx`
    # points to an existing DOCX, use it as a base then inject header/footer.
    def make_reference_with_header_footer(base_ref, header_text):
        # Load existing template or create a new document
        doc = Document(base_ref) if base_ref and os.path.exists(base_ref) else Document()

        def add_page_number_field_to_paragraph(paragraph):
            # Append a simple field for PAGE to the paragraph XML
            fld = OxmlElement('w:fldSimple')
            fld.set(qn('w:instr'), 'PAGE')
            paragraph._p.append(fld)

        for section in doc.sections:
            # Header: center the provided header_text
            header = section.header
            # Use first paragraph or add one
            if header.paragraphs:
                hdr_p = header.paragraphs[0]
                hdr_p.text = header_text
            else:
                hdr_p = header.add_paragraph(header_text)
            hdr_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Footer: center a page-number field
            footer = section.footer
            if footer.paragraphs:
                ftr_p = footer.paragraphs[0]
                # clear any existing text
                ftr_p.text = ''
            else:
                ftr_p = footer.add_paragraph()
            add_page_number_field_to_paragraph(ftr_p)
            ftr_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save to a temp file and return its path
        tmp_ref = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp_path = tmp_ref.name
        tmp_ref.close()
        doc.save(tmp_path)
        return tmp_path

    # Use pypandoc to convert, set resource_path for images, enable raw_tex for page breaks
    # Choose your Pygments theme here (e.g., monokai, tango, zenburn, haddock, etc.)
    pygments_theme = "tango"  # Use a supported theme for DOCX
    pandoc_args = [
        "-f", "markdown+raw_tex",
        f"--syntax-highlighting={pygments_theme}",
        "--filter", "pandoc-pagebreak.py",
        "--lua-filter","style-map.lua"
    ]

    # Always generate a reference DOCX that ensures the header and footer we want.
    # Use any provided reference_docx as a base template (it will be copied and
    # the header/footer injected), otherwise create a minimal template.
    header_text = "Test-First Copilot - Greenfield Edition"
    try:
        generated_ref = make_reference_with_header_footer(reference_docx, header_text)
        pandoc_args.append(f"--reference-doc={generated_ref}")
        # ensure cleanup later
        temp_files.append(generated_ref)
    except Exception as e:
        print("Warning: could not create generated reference docx:", e)
    if folder_name:
        pandoc_args.extend(["--resource-path", folder_name])
    try:
        pypandoc.convert_file(temp_md_file.name, 'docx', outputfile=output_file, extra_args=pandoc_args)
        print(f"Saved {output_file}")
    except Exception as e:
        print("Pandoc error:", e)

    # Clean up temp files
    for tf in temp_files:
        try:
            os.remove(tf)
        except Exception:
            pass

if __name__ == "__main__":
    output_docx = "book.docx"
    # Set reference_docx to your template path or None
    reference_docx = "custom-reference.docx" # None # "template.docx"  # e.g., "reference.docx"
    convert_markdowns_to_docx(output_file=output_docx, chapters_file="chapters.txt", folder_name="Test-First Copilot", reference_docx=reference_docx)
